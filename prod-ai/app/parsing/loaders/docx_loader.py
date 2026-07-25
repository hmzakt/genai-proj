from pathlib import Path
import mimetypes

from docx import Document

from app.parsing.cleaners import TextCleaner
from app.parsing.loaders.base import BaseLoader
from app.parsing.models import (
    ParsedDocument,
    ParsedPage,
    DocumentMetadata,
)
from app.parsing.utils.hashing import sha256_file


class DOCXLoader(BaseLoader):

    def __init__(self):

        self.cleaner = TextCleaner()

    def load(self, file: Path):

        document = Document(file) # type: ignore[arg-type]

        raw = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

        cleaned = self.cleaner.clean(raw)

        page = ParsedPage(
            page_number=1,
            text=cleaned,
            character_count=len(cleaned),
            word_count=len(cleaned.split()),
        )

        metadata = DocumentMetadata(
            filename=file.name,
            extension=file.suffix,
            mime_type=mimetypes.guess_type(file)[0],
            parser_name="DOCXLoader",
            parser_version="1.0",
            page_count=1,
            file_size_bytes=file.stat().st_size,
            sha256=sha256_file(file),
        )

        return ParsedDocument(
            source_path=file,
            extension=file.suffix,
            metadata=metadata,
            pages=[page],
            raw_text=raw,
            cleaned_text=cleaned,
        )